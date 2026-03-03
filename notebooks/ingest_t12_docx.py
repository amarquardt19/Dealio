# Databricks notebook source
# MAGIC %md
# MAGIC # Ingest T12 DOCX → Structured Delta Table
# MAGIC **Deterministic DOCX parsing — no LLM needed for structure detection.**
# MAGIC
# MAGIC Handles 5+ different T12 layouts automatically:
# MAGIC - Month headers: `MM/DD/YYYY`, `Mon YYYY`, `Month YYYY`
# MAGIC - With or without account codes
# MAGIC - Budget/Variance/% Variance columns (detected & extracted)
# MAGIC - Proprietary Labeling column (auto-detected)
# MAGIC - Doc notes (top-of-statement comments)
# MAGIC
# MAGIC **Usage:**
# MAGIC 1. Upload DOCX files to a UC Volume (e.g. `/Volumes/workspace/rawdatalabeling/uploads/`)
# MAGIC 2. Set `DOCX_PATH` and `TABLE_NAME` in the config cell
# MAGIC 3. Run all cells → structured Delta table lands in `rawdatalabeling`
# MAGIC 4. `dbt build` picks it up through `stg_generic_t12`

# COMMAND ----------

# MAGIC %pip install python-docx
# MAGIC dbutils.library.restartPython()

# COMMAND ----------

# MAGIC %md
# MAGIC ## Config

# COMMAND ----------

# --- SET THESE FOR EACH NEW FILE ---
DOCX_PATH = "/Volumes/workspace/rawdatalabeling/uploads/example_t12.docx"
TABLE_NAME = "t12|0125|newproperty"  # format: t12|MMYY|propertyslug

CATALOG = "workspace"
SCHEMA = "rawdatalabeling"

print(f"Input:  {DOCX_PATH}")
print(f"Output: {CATALOG}.{SCHEMA}.`{TABLE_NAME}`")

# COMMAND ----------

# MAGIC %md
# MAGIC ## Core Parsing Engine
# MAGIC Deterministic DOCX table parser — adapted from `t12_docx_cleaning_for_dbt_v5.py`

# COMMAND ----------

import re
import json
import calendar
from datetime import datetime, date
from typing import List, Optional, Dict, Any

import pandas as pd
from docx import Document

# ---------------------------------------------------------------------------
# Regex patterns
# ---------------------------------------------------------------------------
ACCOUNT_RE = re.compile(r"^\s*\d{3,4}-\d{3,4}\s*$")
DATE_RE = re.compile(r"^\s*\d{2}/\d{2}/\d{4}\s*$")
MON_YYYY_RE = re.compile(
    r"^\s*(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)\s+\d{4}\s*$",
    re.IGNORECASE,
)
MON_FULL_YYYY_RE = re.compile(
    r"^\s*(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\s*$",
    re.IGNORECASE,
)
CURRENCY_CLEAN_RE = re.compile(r"[\$,]")

MON_MAP = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
    "jul": 7, "aug": 8, "sep": 9, "sept": 9, "oct": 10, "nov": 11, "dec": 12,
}
MON_FULL_MAP = {
    "january": 1, "february": 2, "march": 3, "april": 4, "may": 5, "june": 6,
    "july": 7, "august": 8, "september": 9, "october": 10, "november": 11, "december": 12,
}

MONTH_ABBR = {
    1: "jan", 2: "feb", 3: "mar", 4: "apr", 5: "may", 6: "jun",
    7: "jul", 8: "aug", 9: "sep", 10: "oct", 11: "nov", 12: "dec",
}

# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------
def cell_text(cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()

def row_texts(row) -> List[str]:
    return [cell_text(c) for c in row.cells]

def norm(s: str) -> str:
    return (s or "").strip().lower()

def is_blank_row(texts):
    return all(norm(t) == "" for t in texts)

def is_spacer_row(texts):
    return all(re.fullmatch(r"[\s,]*", (t or "")) is not None for t in texts)

def is_month_scaffold_row(texts):
    lowered = [norm(t) for t in texts]
    if sum(1 for t in lowered if t == "month ending") >= 3:
        return True
    if sum(1 for t in lowered if t == "actual") >= 3:
        return True
    return "month ending" in " ".join(lowered)

def is_metadata_row(texts):
    joined = " ".join(norm(t) for t in texts)
    keys = [
        "amounts in", "property =", "statement", "period =",
        "reporting book", "as of date", "location",
        "income statement", "trailing income statement",
        "twelve month", "12 month", "book =", "tree =",
    ]
    return any(k in joined for k in keys)

def has_account_code(texts):
    return any(ACCOUNT_RE.match((t or "").strip()) for t in texts[:2])

def extract_account_code(texts):
    for t in texts[:2]:
        tt = (t or "").strip()
        if ACCOUNT_RE.match(tt):
            return tt
    return None

def is_total_row(texts):
    joined = " ".join(norm(t) for t in texts)
    return ("total" in joined) and not has_account_code(texts)

def parse_money(value):
    if value is None:
        return None
    s = str(value).strip()
    if not s or s == "-" or s.upper() == "N/A":
        return None
    s = CURRENCY_CLEAN_RE.sub("", s)
    neg = False
    if s.startswith("(") and s.endswith(")"):
        neg = True
        s = s[1:-1]
    s = s.replace('"', "").strip()
    if not re.fullmatch(r"-?\d+(\.\d+)?", s):
        return None
    x = float(s)
    return -x if neg else x

def is_month_token(t: str) -> bool:
    tt = (t or "").strip()
    return bool(DATE_RE.match(tt) or MON_YYYY_RE.match(tt) or MON_FULL_YYYY_RE.match(tt))

def month_token_to_month_end(t: str) -> Optional[str]:
    tt = (t or "").strip()
    if DATE_RE.match(tt):
        try:
            return datetime.strptime(tt, "%m/%d/%Y").date().isoformat()
        except Exception:
            return None
    if MON_YYYY_RE.match(tt):
        parts = tt.split()
        mon = MON_MAP[parts[0].lower()]
        yr = int(parts[1])
        last_day = calendar.monthrange(yr, mon)[1]
        return date(yr, mon, last_day).isoformat()
    if MON_FULL_YYYY_RE.match(tt):
        parts = tt.split()
        mon = MON_FULL_MAP[parts[0].lower()]
        yr = int(parts[1])
        last_day = calendar.monthrange(yr, mon)[1]
        return date(yr, mon, last_day).isoformat()
    return None

def count_month_tokens(texts: List[str]) -> int:
    return sum(1 for t in texts if is_month_token(t))

# ---------------------------------------------------------------------------
# Table & column detection
# ---------------------------------------------------------------------------
def find_best_table(doc):
    best = None
    best_score = -1
    for t in doc.tables:
        month_strength = 0
        for r in t.rows:
            month_strength = max(month_strength, count_month_tokens(row_texts(r)))
        score = len(t.rows) + (1000 if month_strength >= 6 else 0) + month_strength
        if score > best_score:
            best_score = score
            best = t
    return best

def find_month_row_idx(table) -> Optional[int]:
    best_i = None
    best_c = 0
    for i, r in enumerate(table.rows):
        c = count_month_tokens(row_texts(r))
        if c > best_c:
            best_c = c
            best_i = i
    return best_i if best_c >= 6 else None

def find_proprietary_column(table) -> Optional[int]:
    for r in table.rows:
        texts = row_texts(r)
        for idx, t in enumerate(texts):
            if norm(t) == "proprietary labeling":
                return idx
    return None

def find_keyword_columns(header_texts: List[str]) -> Dict[str, List[int]]:
    out = {"budget": [], "variance": [], "pct_variance": []}
    for idx, t in enumerate(header_texts):
        nt = norm(t)
        if "budget" in nt:
            out["budget"].append(idx)
        if nt == "variance" or ("variance" in nt and "%variance" not in nt):
            out["variance"].append(idx)
        if "%variance" in nt or ("pct" in nt and "variance" in nt):
            out["pct_variance"].append(idx)
    return out

def find_total_column_near_month_row(table, month_row_idx: int, last_month_idx: int) -> int:
    for k in range(1, 6):
        j = month_row_idx - k
        if j < 0:
            break
        texts = row_texts(table.rows[j])
        for idx, t in enumerate(texts):
            if norm(t) == "total":
                return idx
    return last_month_idx + 1

def build_column_map(table) -> Dict[str, Any]:
    month_row_idx = find_month_row_idx(table)
    if month_row_idx is None:
        return {"month_row_idx": None}

    month_row_texts = row_texts(table.rows[month_row_idx])

    month_cols: List[int] = []
    month_ends: List[str] = []
    for idx, t in enumerate(month_row_texts):
        if is_month_token(t):
            me = month_token_to_month_end(t)
            if me:
                month_cols.append(idx)
                month_ends.append(me)

    if not month_cols:
        return {"month_row_idx": month_row_idx}

    last_month_idx = max(month_cols)

    total_col = None
    for idx, t in enumerate(month_row_texts):
        if norm(t) == "total":
            total_col = idx
            break
    if total_col is None:
        total_col = find_total_column_near_month_row(table, month_row_idx, last_month_idx)

    kw = find_keyword_columns(month_row_texts)
    budget_cols = kw["budget"]
    variance_cols = kw["variance"]
    pct_var_cols = kw["pct_variance"]

    return {
        "month_row_idx": month_row_idx,
        "month_cols": month_cols,
        "month_ends": month_ends,
        "total_col": total_col,
        "budget_total_col": budget_cols[0] if len(budget_cols) >= 1 else None,
        "original_budget_total_col": budget_cols[1] if len(budget_cols) >= 2 else None,
        "variance_total_col": variance_cols[0] if variance_cols else None,
        "pct_variance_col": pct_var_cols[0] if pct_var_cols else None,
        "proprietary_col": find_proprietary_column(table),
    }

def has_numeric_payload_by_columns(texts: List[str], numeric_cols: List[int]) -> bool:
    for idx in numeric_cols:
        if idx < len(texts) and parse_money(texts[idx]) is not None:
            return True
    return False

# ---------------------------------------------------------------------------
# Row classification
# ---------------------------------------------------------------------------
def classify_row(texts: List[str], colmap: Dict[str, Any]) -> str:
    if is_blank_row(texts) or is_spacer_row(texts):
        return "spacer"
    if is_metadata_row(texts):
        return "metadata"
    if is_month_scaffold_row(texts):
        return "scaffold"
    if colmap.get("month_row_idx") is not None and count_month_tokens(texts) >= 6:
        return "month_header"

    numeric_cols = (colmap.get("month_cols") or []).copy()
    for k in ["total_col", "budget_total_col", "original_budget_total_col",
              "variance_total_col", "pct_variance_col"]:
        v = colmap.get(k)
        if isinstance(v, int):
            numeric_cols.append(v)

    if not has_account_code(texts) and numeric_cols and has_numeric_payload_by_columns(texts, numeric_cols):
        return "total" if is_total_row(texts) else "subtotal"

    if has_account_code(texts):
        if numeric_cols and has_numeric_payload_by_columns(texts, numeric_cols):
            return "line_item"
        return "header_with_code"

    if is_total_row(texts):
        return "total"
    if numeric_cols and has_numeric_payload_by_columns(texts, numeric_cols):
        return "line_item_no_code"
    if any(norm(t) for t in texts):
        return "header"
    return "unknown"

# ---------------------------------------------------------------------------
# Doc notes extraction
# ---------------------------------------------------------------------------
def extract_doc_notes(table, colmap: Dict[str, Any]) -> List[str]:
    month_row_idx = colmap.get("month_row_idx")
    if month_row_idx is None:
        return []
    numeric_cols = (colmap.get("month_cols") or []).copy()
    for k in ["total_col", "budget_total_col", "original_budget_total_col",
              "variance_total_col", "pct_variance_col"]:
        v = colmap.get(k)
        if isinstance(v, int):
            numeric_cols.append(v)
    notes: List[str] = []
    for i in range(0, month_row_idx):
        texts = row_texts(table.rows[i])
        if is_blank_row(texts) or is_spacer_row(texts) or is_month_scaffold_row(texts) or is_metadata_row(texts):
            continue
        if count_month_tokens(texts) > 0:
            continue
        if has_account_code(texts):
            continue
        if numeric_cols and has_numeric_payload_by_columns(texts, numeric_cols):
            continue
        line = " | ".join([t.strip() for t in texts if (t or "").strip()])
        if line:
            notes.append(line)
    seen = set()
    out = []
    for n in notes:
        if n not in seen:
            seen.add(n)
            out.append(n)
    return out

print("✓ Parsing engine loaded")

# COMMAND ----------

# MAGIC %md
# MAGIC ## Process DOCX → Structured DataFrame

# COMMAND ----------

def process_docx_to_dataframe(doc_path: str) -> pd.DataFrame:
    """
    Parse a DOCX T12 file and return a structured Pandas DataFrame.
    
    Output schema:
      row_type       - header / line_item / line_item_no_code / total / subtotal / header_with_code
      account_code   - e.g. "4100-0100" (null for non-coded rows)
      account_name   - line item description or section header name
      month_01..12   - monthly dollar amounts (null where absent)
      total          - annual total
      proprietary_label - manager-assigned label (if present in source)
      doc_notes      - top-of-document notes/comments
      row_idx        - original row position (for ordering)
    
    ALL rows are included (headers, line items, totals) so the dbt
    intermediate model can detect section headers for classification.
    """
    doc = Document(doc_path)
    table = find_best_table(doc)
    if table is None:
        raise RuntimeError(f"No tables found in {doc_path}")

    any_code = any(has_account_code(row_texts(r)) for r in table.rows)
    desc_col = 1 if any_code else 0

    colmap = build_column_map(table)
    doc_notes_list = extract_doc_notes(table, colmap)
    doc_notes_joined = " || ".join(doc_notes_list) if doc_notes_list else None

    month_cols = colmap.get("month_cols") or []
    month_ends = colmap.get("month_ends") or []
    total_col = colmap.get("total_col")
    proprietary_col = colmap.get("proprietary_col")

    # Build month column names: month_01..month_12 in ordinal order
    # Also build friendly names like "apr_2022" for metadata
    month_friendly_names = []
    for me in month_ends:
        d = date.fromisoformat(me)
        month_friendly_names.append(f"{MONTH_ABBR[d.month]}_{d.year}")

    keep_types = {"header", "header_with_code", "line_item", "line_item_no_code", "total", "subtotal"}
    output_rows = []

    for i, r in enumerate(table.rows):
        texts = row_texts(r)
        row_type = classify_row(texts, colmap)
        if row_type not in keep_types:
            continue

        account_code = extract_account_code(texts) if any_code else None
        account_name = texts[desc_col].strip() if desc_col < len(texts) else ""

        # For header rows that have no description in desc_col, try all cells
        if not account_name and row_type in ("header", "header_with_code"):
            account_name = " ".join(t.strip() for t in texts if t.strip()).strip()

        proprietary_label = None
        if isinstance(proprietary_col, int) and proprietary_col < len(texts):
            proprietary_label = texts[proprietary_col].strip() or None

        total_val = parse_money(texts[total_col]) if isinstance(total_col, int) and total_col < len(texts) else None

        row_data = {
            "row_type": row_type,
            "account_code": account_code,
            "account_name": account_name,
            "total": total_val,
            "proprietary_label": proprietary_label,
            "doc_notes": doc_notes_joined,
            "row_idx": i,
        }

        # Extract monthly values
        for m_idx, col_idx in enumerate(month_cols):
            col_name = f"month_{str(m_idx + 1).zfill(2)}"
            row_data[col_name] = parse_money(texts[col_idx]) if col_idx < len(texts) else None

        output_rows.append(row_data)

    df = pd.DataFrame(output_rows)

    # Ensure all month columns exist (month_01..month_12) even if fewer months detected
    for m in range(1, 13):
        col_name = f"month_{str(m).zfill(2)}"
        if col_name not in df.columns:
            df[col_name] = None

    # Reorder columns
    ordered_cols = ["row_type", "account_code", "account_name"]
    ordered_cols += [f"month_{str(m).zfill(2)}" for m in range(1, 13)]
    ordered_cols += ["total", "proprietary_label", "doc_notes", "row_idx"]
    df = df[[c for c in ordered_cols if c in df.columns]]

    # Print summary
    print(f"✓ Parsed {len(df)} rows from {doc_path}")
    print(f"  Row types: {df['row_type'].value_counts().to_dict()}")
    print(f"  Months detected: {len(month_ends)}")
    if month_friendly_names:
        print(f"  Month range: {month_friendly_names[0]} → {month_friendly_names[-1]}")
    print(f"  Column map: {json.dumps({k: v for k, v in colmap.items() if k != 'month_row_idx'}, default=str)}")
    if doc_notes_list:
        print(f"  Doc notes: {doc_notes_list}")

    return df, colmap, month_friendly_names

# COMMAND ----------

# MAGIC %md
# MAGIC ## Run: Parse DOCX & Preview

# COMMAND ----------

pdf, colmap, month_names = process_docx_to_dataframe(DOCX_PATH)
display(spark.createDataFrame(pdf))

# COMMAND ----------

# MAGIC %md
# MAGIC ## Write to Delta Table
# MAGIC Writes the structured output to `rawdatalabeling` so the dbt pipeline picks it up.
# MAGIC
# MAGIC The table includes header rows so `int_generic_t12` can detect section boundaries.

# COMMAND ----------

sdf = spark.createDataFrame(pdf)

full_table_name = f"{CATALOG}.{SCHEMA}.`{TABLE_NAME}`"
sdf.write.mode("overwrite").saveAsTable(full_table_name)

row_count = sdf.count()
print(f"✓ Written {row_count} rows to {full_table_name}")
print(f"\nNext step: add '{TABLE_NAME}' to sources.yml and stg_generic_t12 source list, then run:")
print(f"  dbt build")

# COMMAND ----------

# MAGIC %md
# MAGIC ## Verify: Preview the Delta Table

# COMMAND ----------

display(spark.sql(f"SELECT * FROM {CATALOG}.{SCHEMA}.`{TABLE_NAME}` ORDER BY row_idx"))

# COMMAND ----------

# MAGIC %md
# MAGIC ## Metadata: Save Column Map & Month Names
# MAGIC Stored alongside the table for reference.

# COMMAND ----------

metadata = {
    "table_name": TABLE_NAME,
    "docx_path": DOCX_PATH,
    "month_names": month_names,
    "column_map": {k: str(v) for k, v in colmap.items()},
    "row_count": int(pdf.shape[0]),
    "source": "ingest_t12_docx notebook",
}

spark.sql(f"""
    COMMENT ON TABLE {CATALOG}.{SCHEMA}.`{TABLE_NAME}`
    IS '{json.dumps(metadata).replace("'", "''")}'
""")
print(f"✓ Metadata saved as table comment")
print(json.dumps(metadata, indent=2))
