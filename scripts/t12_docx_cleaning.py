#!/usr/bin/env python3
"""
t12_docx_cleaning_for_dbt_v5.py

Deterministic DOCX -> staging tables for dbt (no LLM).

New upgrade for "The Archer" style:
- Captures top-of-statement COMMENT / NOTE rows (e.g., "Taxes on a 299c assessment freeze...")
  and writes them to <out_prefix>_doc_notes.json + <out_prefix>_doc_notes.txt
  and also attaches `doc_notes` (joined string) to every output record.

Still supports:
- Month header detection: MM/DD/YYYY OR Mon YYYY (short/full month names)
- Column mapping by header row (months, Total, Budget/Variance/%Variance if present, Proprietary Labeling)
- Extract month values by column index
- "N/A" -> NULL
- Account codes: 3–4 digits before dash, 3–4 after dash
- Totals/subtotals: (no code + numeric payload) => subtotal/total
- Coded headers without numeric payload => header_with_code

Outputs (CSV):
  - <out_prefix>_raw_rows.csv
  - <out_prefix>_canonical_rows.csv
  - <out_prefix>_line_item_month.csv
  - <out_prefix>_line_item_wide.csv
  - <out_prefix>_colmap.json
  - <out_prefix>_doc_notes.json
  - <out_prefix>_doc_notes.txt
"""

import argparse
import re
import json
import calendar
from datetime import datetime, date
from typing import List, Optional, Dict, Any

import pandas as pd
from docx import Document

ACCOUNT_RE = re.compile(r"^\s*\d{3,4}-\d{3,4}\s*$")
DATE_RE = re.compile(r"^\s*\d{2}/\d{2}/\d{4}\s*$")
MON_YYYY_RE = re.compile(r"^\s*(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)\s+\d{4}\s*$", re.IGNORECASE)
MON_FULL_YYYY_RE = re.compile(
    r"^\s*(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\s*$",
    re.IGNORECASE
)
CURRENCY_CLEAN_RE = re.compile(r"[\$,]")

MON_MAP = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
    "jul": 7, "aug": 8, "sep": 9, "sept": 9, "oct": 10, "nov": 11, "dec": 12
}
MON_FULL_MAP = {
    "january": 1, "february": 2, "march": 3, "april": 4, "may": 5, "june": 6,
    "july": 7, "august": 8, "september": 9, "october": 10, "november": 11, "december": 12
}

def cell_text(cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()

def row_texts(row) -> List[str]:
    return [cell_text(c) for c in row.cells]

def norm(s: str) -> str:
    return (s or "").strip().lower()

def is_blank_row(texts):
    return all(norm(t) == "" for t in texts)

def is_spacer_row(texts):
    return all(re.fullmatch(r"[\s,]*", (t or "")) is not None for t in texts)

def is_month_scaffold_row(texts):
    lowered = [norm(t) for t in texts]
    if sum(1 for t in lowered if t == "month ending") >= 3:
        return True
    if sum(1 for t in lowered if t == "actual") >= 3:
        return True
    return "month ending" in " ".join(lowered)

def is_metadata_row(texts):
    joined = " ".join(norm(t) for t in texts)
    keys = [
        "amounts in", "property =", "statement", "period =",
        "reporting book", "as of date", "location",
        "income statement", "trailing income statement", "twelve month", "12 month",
        "book =", "tree ="
    ]
    return any(k in joined for k in keys)

def has_account_code(texts):
    return any(ACCOUNT_RE.match((t or "").strip()) for t in texts[:2])

def extract_account_code(texts):
    for t in texts[:2]:
        tt = (t or "").strip()
        if ACCOUNT_RE.match(tt):
            return tt
    return None

def is_total_row(texts):
    joined = " ".join(norm(t) for t in texts)
    return ("total" in joined) and not has_account_code(texts)

def parse_money(value):
    if value is None:
        return None
    s = str(value).strip()
    if not s or s == "-" or s.upper() == "N/A":
        return None
    s = CURRENCY_CLEAN_RE.sub("", s)
    neg = False
    if s.startswith("(") and s.endswith(")"):
        neg = True
        s = s[1:-1]
    s = s.replace('"', '').strip()
    if not re.fullmatch(r"-?\d+(\.\d+)?", s):
        return None
    x = float(s)
    return -x if neg else x

def is_month_token(t: str) -> bool:
    tt = (t or "").strip()
    return bool(DATE_RE.match(tt) or MON_YYYY_RE.match(tt) or MON_FULL_YYYY_RE.match(tt))

def month_token_to_month_end(t: str) -> Optional[str]:
    tt = (t or "").strip()
    if DATE_RE.match(tt):
        try:
            return datetime.strptime(tt, "%m/%d/%Y").date().isoformat()
        except Exception:
            return None
    if MON_YYYY_RE.match(tt):
        parts = tt.split()
        mon = MON_MAP[parts[0].lower()]
        yr = int(parts[1])
        last_day = calendar.monthrange(yr, mon)[1]
        return date(yr, mon, last_day).isoformat()
    if MON_FULL_YYYY_RE.match(tt):
        parts = tt.split()
        mon = MON_FULL_MAP[parts[0].lower()]
        yr = int(parts[1])
        last_day = calendar.monthrange(yr, mon)[1]
        return date(yr, mon, last_day).isoformat()
    return None

def count_month_tokens(texts: List[str]) -> int:
    return sum(1 for t in texts if is_month_token(t))

def find_best_table(doc):
    best = None
    best_score = -1
    for t in doc.tables:
        month_strength = 0
        for r in t.rows:
            month_strength = max(month_strength, count_month_tokens(row_texts(r)))
        score = len(t.rows) + (1000 if month_strength >= 6 else 0) + month_strength
        if score > best_score:
            best_score = score
            best = t
    return best

def find_month_row_idx(table) -> Optional[int]:
    best_i = None
    best_c = 0
    for i, r in enumerate(table.rows):
        c = count_month_tokens(row_texts(r))
        if c > best_c:
            best_c = c
            best_i = i
    return best_i if best_c >= 6 else None

def find_proprietary_column(table) -> Optional[int]:
    for r in table.rows:
        texts = row_texts(r)
        for idx, t in enumerate(texts):
            if norm(t) == "proprietary labeling":
                return idx
    return None

def find_keyword_columns(header_texts: List[str]) -> Dict[str, List[int]]:
    out = {"budget": [], "variance": [], "pct_variance": []}
    for idx, t in enumerate(header_texts):
        nt = norm(t)
        if "budget" in nt:
            out["budget"].append(idx)
        if nt == "variance" or ("variance" in nt and "%variance" not in nt):
            out["variance"].append(idx)
        if "%variance" in nt or ("pct" in nt and "variance" in nt):
            out["pct_variance"].append(idx)
    return out

def find_total_column_near_month_row(table, month_row_idx: int, last_month_idx: int) -> int:
    for k in range(1, 6):
        j = month_row_idx - k
        if j < 0:
            break
        texts = row_texts(table.rows[j])
        for idx, t in enumerate(texts):
            if norm(t) == "total":
                return idx
    return last_month_idx + 1

def build_column_map(table) -> Dict[str, Any]:
    month_row_idx = find_month_row_idx(table)
    if month_row_idx is None:
        return {"month_row_idx": None}

    month_row_texts = row_texts(table.rows[month_row_idx])

    month_cols: List[int] = []
    month_ends: List[str] = []
    for idx, t in enumerate(month_row_texts):
        if is_month_token(t):
            me = month_token_to_month_end(t)
            if me:
                month_cols.append(idx)
                month_ends.append(me)

    if not month_cols:
        return {"month_row_idx": month_row_idx}

    last_month_idx = max(month_cols)

    total_col = None
    for idx, t in enumerate(month_row_texts):
        if norm(t) == "total":
            total_col = idx
            break
    if total_col is None:
        total_col = find_total_column_near_month_row(table, month_row_idx, last_month_idx)

    kw = find_keyword_columns(month_row_texts)
    budget_cols = kw["budget"]
    variance_cols = kw["variance"]
    pct_var_cols = kw["pct_variance"]

    budget_total_col = budget_cols[0] if len(budget_cols) >= 1 else None
    original_budget_total_col = budget_cols[1] if len(budget_cols) >= 2 else None
    variance_total_col = variance_cols[0] if variance_cols else None
    pct_variance_col = pct_var_cols[0] if pct_var_cols else None

    proprietary_col = find_proprietary_column(table)

    return {
        "month_row_idx": month_row_idx,
        "month_cols": month_cols,
        "month_ends": month_ends,
        "total_col": total_col,
        "budget_total_col": budget_total_col,
        "original_budget_total_col": original_budget_total_col,
        "variance_total_col": variance_total_col,
        "pct_variance_col": pct_variance_col,
        "proprietary_col": proprietary_col,
    }

def has_numeric_payload_by_columns(texts: List[str], numeric_cols: List[int]) -> bool:
    for idx in numeric_cols:
        if idx < len(texts) and parse_money(texts[idx]) is not None:
            return True
    return False

def extract_doc_notes(table, colmap: Dict[str, Any]) -> List[str]:
    month_row_idx = colmap.get("month_row_idx")
    if month_row_idx is None:
        return []

    numeric_cols = (colmap.get("month_cols") or []).copy()
    for k in ["total_col", "budget_total_col", "original_budget_total_col", "variance_total_col", "pct_variance_col"]:
        v = colmap.get(k)
        if isinstance(v, int):
            numeric_cols.append(v)

    notes: List[str] = []
    for i in range(0, month_row_idx):
        texts = row_texts(table.rows[i])
        if is_blank_row(texts) or is_spacer_row(texts) or is_month_scaffold_row(texts) or is_metadata_row(texts):
            continue
        if count_month_tokens(texts) > 0:
            continue
        if has_account_code(texts):
            continue
        if numeric_cols and has_numeric_payload_by_columns(texts, numeric_cols):
            continue
        line = " | ".join([t.strip() for t in texts if (t or "").strip()])
        if line:
            notes.append(line)

    seen = set()
    out = []
    for n in notes:
        if n not in seen:
            seen.add(n)
            out.append(n)
    return out

def classify_row(texts: List[str], colmap: Dict[str, Any]) -> str:
    if is_blank_row(texts) or is_spacer_row(texts):
        return "spacer"
    if is_metadata_row(texts):
        return "metadata"
    if is_month_scaffold_row(texts):
        return "scaffold"

    if colmap.get("month_row_idx") is not None and count_month_tokens(texts) >= 6:
        return "month_header"

    numeric_cols = (colmap.get("month_cols") or []).copy()
    for k in ["total_col", "budget_total_col", "original_budget_total_col", "variance_total_col", "pct_variance_col"]:
        v = colmap.get(k)
        if isinstance(v, int):
            numeric_cols.append(v)

    if not has_account_code(texts) and numeric_cols and has_numeric_payload_by_columns(texts, numeric_cols):
        return "total" if is_total_row(texts) else "subtotal"

    if has_account_code(texts):
        if numeric_cols and has_numeric_payload_by_columns(texts, numeric_cols):
            return "line_item"
        return "header_with_code"

    if is_total_row(texts):
        return "total"

    if numeric_cols and has_numeric_payload_by_columns(texts, numeric_cols):
        return "line_item_no_code"

    if any(norm(t) for t in texts):
        return "header"
    return "unknown"

def run(doc_path: str, out_prefix: str):
    doc = Document(doc_path)
    table = find_best_table(doc)
    if table is None:
        raise RuntimeError("No tables found in DOCX.")

    any_code = any(has_account_code(row_texts(r)) for r in table.rows)
    desc_col = 1 if any_code else 0

    colmap = build_column_map(table)
    doc_notes_list = extract_doc_notes(table, colmap)
    doc_notes_joined = " || ".join(doc_notes_list) if doc_notes_list else None

    with open(f"{out_prefix}_doc_notes.json", "w", encoding="utf-8") as f:
        json.dump({"doc_path": doc_path, "doc_notes": doc_notes_list}, f, indent=2)
    with open(f"{out_prefix}_doc_notes.txt", "w", encoding="utf-8") as f:
        if doc_notes_list:
            for n in doc_notes_list:
                f.write(n + "\n")

    rows: List[Dict[str, Any]] = []
    for i, r in enumerate(table.rows):
        texts = row_texts(r)
        row_type = classify_row(texts, colmap)
        rows.append({"row_idx": i, "row_type": row_type, "cells": texts})

    pd.DataFrame(rows).to_csv(f"{out_prefix}_raw_rows.csv", index=False)

    drop_types = {"spacer", "scaffold", "header", "header_with_code", "unknown", "month_header", "metadata"}
    canonical_df = pd.DataFrame(rows)
    canonical_df = canonical_df[~canonical_df["row_type"].isin(drop_types)].copy()
    canonical_df.to_csv(f"{out_prefix}_canonical_rows.csv", index=False)

    month_cols = colmap.get("month_cols") or []
    month_ends = colmap.get("month_ends") or []
    total_col = colmap.get("total_col")
    budget_total_col = colmap.get("budget_total_col")
    original_budget_total_col = colmap.get("original_budget_total_col")
    variance_total_col = colmap.get("variance_total_col")
    pct_variance_col = colmap.get("pct_variance_col")
    proprietary_col = colmap.get("proprietary_col")

    facts: List[Dict[str, Any]] = []
    wide_rows: List[Dict[str, Any]] = []

    for rec in rows:
        if rec["row_type"] not in {"line_item", "line_item_no_code", "total", "subtotal"}:
            continue
        cells = rec["cells"]

        account_code = extract_account_code(cells) if any_code else None
        line_item = cells[desc_col].strip() if desc_col < len(cells) else ""

        proprietary_label = None
        if isinstance(proprietary_col, int) and proprietary_col < len(cells):
            proprietary_label = cells[proprietary_col].strip() or None

        total_val = parse_money(cells[total_col]) if isinstance(total_col, int) and total_col < len(cells) else None
        budget_total_val = parse_money(cells[budget_total_col]) if isinstance(budget_total_col, int) and budget_total_col < len(cells) else None
        original_budget_total_val = parse_money(cells[original_budget_total_col]) if isinstance(original_budget_total_col, int) and original_budget_total_col < len(cells) else None
        variance_total_val = parse_money(cells[variance_total_col]) if isinstance(variance_total_col, int) and variance_total_col < len(cells) else None
        pct_variance_val = parse_money(cells[pct_variance_col]) if isinstance(pct_variance_col, int) and pct_variance_col < len(cells) else None

        wide = {
            "doc_path": doc_path,
            "row_idx": rec["row_idx"],
            "row_type": rec["row_type"],
            "account_code": account_code,
            "line_item": line_item,
            "total": total_val,
            "budget_total": budget_total_val,
            "original_budget_total": original_budget_total_val,
            "variance_total": variance_total_val,
            "pct_variance": pct_variance_val,
            "proprietary_label": proprietary_label,
            "doc_notes": doc_notes_joined,
        }

        if month_cols and month_ends and len(month_cols) == len(month_ends):
            for idx, me in zip(month_cols, month_ends):
                amt = parse_money(cells[idx]) if idx < len(cells) else None
                facts.append({
                    "doc_path": doc_path,
                    "row_idx": rec["row_idx"],
                    "row_type": rec["row_type"],
                    "account_code": account_code,
                    "line_item": line_item,
                    "month_end": me,
                    "amount": amt,
                    "total": total_val,
                    "budget_total": budget_total_val,
                    "original_budget_total": original_budget_total_val,
                    "variance_total": variance_total_val,
                    "pct_variance": pct_variance_val,
                    "proprietary_label": proprietary_label,
                    "doc_notes": doc_notes_joined,
                })
                wide[f"amt_{me}"] = amt
        else:
            facts.append({
                "doc_path": doc_path,
                "row_idx": rec["row_idx"],
                "row_type": rec["row_type"],
                "account_code": account_code,
                "line_item": line_item,
                "month_end": None,
                "amount": None,
                "total": total_val,
                "budget_total": budget_total_val,
                "original_budget_total": original_budget_total_val,
                "variance_total": variance_total_val,
                "pct_variance": pct_variance_val,
                "proprietary_label": proprietary_label,
                "doc_notes": doc_notes_joined,
            })

        wide_rows.append(wide)

    pd.DataFrame(facts).to_csv(f"{out_prefix}_line_item_month.csv", index=False)
    pd.DataFrame(wide_rows).to_csv(f"{out_prefix}_line_item_wide.csv", index=False)
    pd.Series(colmap).to_json(f"{out_prefix}_colmap.json", indent=2)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--doc", required=True, help="Path to DOCX file")
    ap.add_argument("--out_prefix", required=True, help="Output prefix (folder + base name), without extension")
    args = ap.parse_args()
    run(args.doc, args.out_prefix)

if __name__ == "__main__":
    main()
